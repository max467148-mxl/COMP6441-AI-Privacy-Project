import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Report-z5557885-final.md"
SUBMISSION = ROOT / "submission"
OUT = SUBMISSION / "Report-z5557885.docx"
ASSETS = ROOT / "docs" / "generated_assets"

BLUE = "235789"
NAVY = "17324D"
CYAN = "2A9D8F"
LIGHT = "EDF3F7"
TEXT = "202A33"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    paragraph.add_run()._r.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run("1")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 12.5, BLUE, 10, 4),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "COMP6441  |  AI Context Aggregation and Privacy Leakage"
    header.style = doc.styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(section.footer.paragraphs[0])


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    available = 6.74
    weights = [1.1] * len(rows[0])
    if len(rows[0]) == 3:
        weights = [2.2, 1.2, 3.3]
    elif len(rows[0]) == 4:
        weights = [0.9, 2.3, 1.0, 3.2] if rows[0][0].strip() == "Date" else [2.7, 0.7, 1.1, 1.4]
    total = sum(weights)
    for r_index, values in enumerate(rows):
        cells = table.rows[0].cells if r_index == 0 else table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Inches(available * weights[i] / total)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value.strip())
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if r_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if r_index == 0:
                shade(cells[i], BLUE)
            elif r_index % 2 == 0:
                shade(cells[i], "F4F7F9")
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def font(size, bold=False):
    candidates = ["C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def build_threat_model():
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / "threat_model.png"
    image = Image.new("RGB", (1600, 520), "white")
    draw = ImageDraw.Draw(image)
    labels = [
        ("Synthetic\nfragments", "Collection"),
        ("Context\nselection", "Minimisation boundary"),
        ("Prompt +\nAI model", "Inference boundary"),
        ("Sensitive\nprofile", "Disclosure surface"),
    ]
    colors = ["#EDF3F7", "#DDF1EC", "#E8EEF7", "#F9E3DE"]
    xs = [70, 445, 820, 1195]
    for index, ((label, caption), x) in enumerate(zip(labels, xs)):
        draw.rounded_rectangle((x, 135, x + 260, 340), radius=18, fill=colors[index], outline="#235789", width=4)
        lines = label.split("\n")
        y = 175
        for line in lines:
            box = draw.textbbox((0, 0), line, font=font(34, True))
            draw.text((x + 130 - (box[2] - box[0]) / 2, y), line, fill="#17324D", font=font(34, True))
            y += 44
        box = draw.textbbox((0, 0), caption, font=font(20))
        draw.text((x + 130 - (box[2] - box[0]) / 2, 292), caption, fill="#4E6575", font=font(20))
        if index < 3:
            draw.line((x + 275, 238, xs[index + 1] - 20, 238), fill="#2A9D8F", width=8)
            draw.polygon([(xs[index + 1] - 20, 222), (xs[index + 1], 238), (xs[index + 1] - 20, 254)], fill="#2A9D8F")
    draw.text((70, 55), "Privacy risk accumulates when unrelated context crosses boundaries", fill="#235789", font=font(38, True))
    draw.text((70, 415), "Primary threats: linkability  |  attribute inference  |  disclosure  |  user unawareness", fill="#344B5C", font=font(26))
    image.save(path)
    return path


def normalise_chart(source, name):
    """Strip alpha metadata that can break headless office PDF export."""
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    with Image.open(source) as image:
        image.convert("RGB").save(path)
    return path


def add_figure(doc, key, number):
    paths = {
        "threat_model": build_threat_model(),
        "condition": normalise_chart(ROOT / "results" / "formal_analysis" / "leakage_by_condition.png", "leakage_by_condition.png"),
        "category": normalise_chart(ROOT / "results" / "formal_analysis" / "leakage_by_category.png", "leakage_by_category.png"),
        "mitigation": normalise_chart(ROOT / "results" / "formal_analysis" / "mitigation_comparison.png", "mitigation_comparison.png"),
    }
    captions = {
        "threat_model": "Threat model and privacy-relevant trust boundaries.",
        "condition": "Mean leakage score by baseline context condition (n = 15 per condition).",
        "category": "Mean baseline leakage score by inference category (n = 12 per category).",
        "mitigation": "Mitigation comparison under full aggregated context (n = 15 per treatment).",
    }
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.add_run().add_picture(str(paths[key]), width=Inches(6.35))
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(f"Figure {number}. {captions[key]}")


def add_cover(doc):
    doc.add_paragraph("\n\n")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("COMP6441  |  CYBERSECURITY INDEPENDENT PROJECT")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(CYAN)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(32)
    run = title.add_run("When Harmless Fragments\nBecome Sensitive")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Measuring Privacy Leakage Through AI Context Aggregation")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph("\n")
    rule = doc.add_table(rows=1, cols=1)
    rule.cell(0, 0).height = Inches(0.08)
    shade(rule.cell(0, 0), CYAN)
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.paragraph_format.space_before = Pt(36)
    add_inline(details, "**Xiaolong Ma**\nz5557885\n22 July 2026")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(44)
    run = note.add_run("90 controlled prompts  |  3 synthetic profiles  |  4 context conditions  |  2 mitigations")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("526A78")
    doc.add_page_break()


def build():
    SUBMISSION.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    add_cover(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line == "## Abstract")
    figure_number = 0
    i = start
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line:
            i += 1
            continue
        marker = re.fullmatch(r"\[\[FIGURE:(.+)\]\]", line)
        if marker:
            figure_number += 1
            add_figure(doc, marker.group(1), figure_number)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            rows = [rows[0]] + rows[2:]
            add_table(doc, rows)
            continue
        if line.startswith("### "):
            title = line[4:]
            if title.startswith("3.3 Original proposal"):
                doc.add_page_break()
            doc.add_heading(title, level=2)
        elif line.startswith("## "):
            title = line[3:]
            if title in {"Appendix A. Evidence Map"}:
                doc.add_page_break()
            doc.add_heading(title, level=1)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            add_inline(p, line)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            paragraph_lines = [line]
            while i + 1 < len(lines):
                nxt = lines[i + 1].rstrip()
                if not nxt or nxt.startswith(("#", "|", "- ", "```", "[[FIGURE:")) or re.match(r"^\d+\. ", nxt):
                    break
                paragraph_lines.append(nxt)
                i += 1
            p = doc.add_paragraph()
            add_inline(p, " ".join(paragraph_lines))
        i += 1

    properties = doc.core_properties
    properties.title = "When Harmless Fragments Become Sensitive"
    properties.subject = "COMP6441 Cybersecurity Independent Project"
    properties.author = "Xiaolong Ma"
    properties.keywords = "privacy, AI, context aggregation, LINDDUN, data minimisation"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
